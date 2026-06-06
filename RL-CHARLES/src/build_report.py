from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt


REPORT_MD = Path("reports/laporan_proyek_utama_informatika.md")
REPORT_DOCX = Path("reports/laporan_proyek_utama_informatika.docx")
METRICS_PATH = Path("reports/model_metrics.json")
IMPORTANCE_PATH = Path("reports/feature_importance.csv")
DATA_PATH = Path("data/dummy_penjualan_umkm.csv")


def rupiah(value: float) -> str:
    return "Rp {:,.0f}".format(value).replace(",", ".")


def load_context() -> dict[str, object]:
    df = pd.read_csv(DATA_PATH, parse_dates=["tanggal"])
    with METRICS_PATH.open("r", encoding="utf-8") as file:
        metrics = json.load(file)
    importance = pd.read_csv(IMPORTANCE_PATH).head(8)
    return {"df": df, "metrics": metrics, "importance": importance}


def build_markdown() -> str:
    context = load_context()
    df: pd.DataFrame = context["df"]
    metrics: dict[str, float] = context["metrics"]
    importance: pd.DataFrame = context["importance"]

    importance_rows = "\n".join(
        f"| {row.feature} | {row.importance:.4f} |" for row in importance.itertuples(index=False)
    )

    return f"""# Sistem Prediksi Penjualan UMKM Menggunakan Random Forest

## Identitas Proyek

Judul proyek: Sistem Prediksi Penjualan UMKM Menggunakan Random Forest

Jenis proyek: Proyek Utama Informatika

Status data: Dataset dummy sintetis untuk prototype akademik

## Abstrak

Proyek ini mengembangkan prototype sistem prediksi penjualan untuk UMKM umum menggunakan algoritma Random Forest Regression. Sistem dibuat untuk membantu pemilik UMKM memperkirakan total penjualan harian berdasarkan faktor produk, harga, stok awal, promo, kanal penjualan, waktu, cuaca, jumlah pengunjung, dan rating produk. Dataset yang digunakan adalah data sintetis yang dirancang menyerupai pola operasional UMKM, sehingga aman digunakan untuk pembelajaran dan laporan akademik tanpa mengekspos data transaksi nyata. Hasil evaluasi pada data uji menunjukkan MAE sebesar {rupiah(metrics['mae'])}, RMSE sebesar {rupiah(metrics['rmse'])}, R2 sebesar {metrics['r2']:.3f}, dan MAPE sebesar {metrics['mape']:.2f}%. Prototype dilengkapi notebook eksperimen, model tersimpan, dan aplikasi Streamlit multipage.

Kata kunci: UMKM, prediksi penjualan, Random Forest, machine learning, prototype.

## BAB I. Pendahuluan

### 1.1 Latar Belakang

UMKM sering menghadapi ketidakpastian penjualan harian. Perubahan harga, stok, promosi, kanal penjualan, hari libur, dan jumlah pengunjung dapat memengaruhi pendapatan. Tanpa alat prediksi sederhana, keputusan stok dan promosi biasanya dibuat berdasarkan intuisi. Kondisi tersebut dapat menyebabkan kelebihan stok, kekurangan stok, atau promosi yang kurang tepat.

Machine learning dapat digunakan untuk mengenali pola historis dan menghasilkan prediksi penjualan. Random Forest dipilih karena mampu memodelkan hubungan non-linear pada data tabular, relatif tahan terhadap outlier, dan dapat memberikan feature importance yang mudah dijelaskan kepada pengguna non-teknis.

### 1.2 Rumusan Masalah

Rumusan masalah proyek ini adalah bagaimana merancang prototype sistem prediksi total penjualan harian UMKM menggunakan Random Forest dengan dataset dummy yang realistis, aman, dan mudah diganti dengan data asli pada tahap lanjutan.

### 1.3 Tujuan

Tujuan proyek ini adalah:

1. Membuat dataset dummy penjualan UMKM yang masuk akal secara bisnis.
2. Membangun model Random Forest Regression untuk memprediksi total penjualan harian.
3. Menyediakan notebook eksperimen yang sistematis dan reproducible.
4. Membuat aplikasi Streamlit multipage untuk demo prediksi dan evaluasi model.
5. Menyusun laporan akademik proyek dengan metodologi, hasil, dan keterbatasan yang jelas.

### 1.4 Batasan Masalah

Prototype ini menggunakan dataset sintetis, bukan data transaksi UMKM nyata. Model hanya memprediksi total penjualan per produk per hari. Sistem belum mencakup integrasi POS, monitoring produksi, retraining otomatis, atau validasi lapangan oleh pemilik UMKM.

## BAB II. Tinjauan Pustaka

### 2.1 Prediksi Penjualan UMKM

Prediksi penjualan adalah proses memperkirakan nilai penjualan pada periode mendatang berdasarkan pola historis dan faktor pendukung. Untuk UMKM, prediksi dapat membantu perencanaan stok, promosi, dan pengambilan keputusan operasional.

### 2.2 Random Forest Regression

Random Forest adalah metode ensemble yang membangun banyak decision tree dan menggabungkan hasilnya untuk meningkatkan stabilitas prediksi. Breiman (2001) menjelaskan Random Forest sebagai kombinasi tree predictor yang memanfaatkan randomization untuk mengurangi variansi. Pada tugas regresi, hasil akhir diperoleh dari rata-rata prediksi banyak pohon.

### 2.3 Penelitian Terkait

Punia et al. (2020) membahas penggunaan LSTM dan Random Forest untuk demand forecasting pada retail multi-channel. Studi tersebut relevan karena menunjukkan bahwa Random Forest masih berguna sebagai model tabular yang kuat untuk konteks permintaan dan penjualan. Pada proyek ini, Random Forest dipakai dalam skala prototype akademik dengan dataset sintetis.

## BAB III. Metodologi

### 3.1 Dataset

Dataset dibuat secara sintetis dengan {len(df):,} baris, {df['id_produk'].nunique()} produk, dan periode {df['tanggal'].min().date()} sampai {df['tanggal'].max().date()}. Data merepresentasikan UMKM umum dengan kategori makanan_minuman, fashion, kebutuhan_rumah, aksesoris, dan kerajinan.

Fitur utama meliputi harga satuan, stok awal, diskon, status promo, kanal penjualan, hari dalam minggu, akhir pekan, bulan, musim Ramadan/Lebaran, cuaca, jumlah pengunjung, rating produk, nama produk, dan kategori produk. Target prediksi adalah total_penjualan.

### 3.2 Keamanan dan Etika Data

Dataset tidak menggunakan nama pelanggan, nomor telepon, alamat, transaksi nyata, atau identitas bisnis asli. Seluruh data dibuat untuk simulasi akademik. Karena itu, dataset aman untuk dipublikasikan pada laporan proyek dan demonstrasi kelas.

### 3.3 Preprocessing

Fitur numerik diproses menggunakan imputasi median dan standardisasi. Fitur kategorikal diproses menggunakan imputasi modus dan One-Hot Encoding. Pipeline scikit-learn digunakan agar preprocessing dan model tersimpan dalam satu artefak yang konsisten.

### 3.4 Pembagian Data

Data dibagi berdasarkan waktu. Data latih menggunakan periode {metrics['train_start']} sampai {metrics['train_end']}, sedangkan data uji menggunakan periode {metrics['test_start']} sampai {metrics['test_end']}. Split berbasis waktu dipilih agar evaluasi lebih mendekati skenario prediksi masa depan dan mengurangi risiko data leakage.

### 3.5 Model

Model utama adalah RandomForestRegressor dengan 300 estimator, max_depth 18, min_samples_leaf 2, random_state 42, dan n_jobs -1. Fitur yang berpotensi bocor seperti jumlah_terjual, stok_akhir, dan biaya_promosi tidak digunakan sebagai input model karena nilainya diketahui setelah transaksi terjadi.

### 3.6 Metrik Evaluasi

Evaluasi menggunakan MAE, RMSE, R2, dan MAPE. MAE dan RMSE menunjukkan besar error dalam rupiah. R2 menunjukkan proporsi variasi target yang dapat dijelaskan model. MAPE menunjukkan error relatif dalam persen.

## BAB IV. Hasil dan Pembahasan

### 4.1 Ringkasan Data

Total penjualan sintetis selama periode observasi adalah {rupiah(df['total_penjualan'].sum())}. Rata-rata total penjualan per baris adalah {rupiah(df['total_penjualan'].mean())}. Dataset menunjukkan pola yang disengaja: promosi meningkatkan penjualan, akhir pekan cenderung lebih tinggi, stok awal membatasi penjualan maksimum, dan momen musiman menaikkan permintaan kategori tertentu.

### 4.2 Hasil Evaluasi Model

| Metrik | Nilai |
|---|---:|
| MAE | {rupiah(metrics['mae'])} |
| RMSE | {rupiah(metrics['rmse'])} |
| R2 | {metrics['r2']:.3f} |
| MAPE | {metrics['mape']:.2f}% |
| CV MAE Mean | {rupiah(metrics['cv_mae_mean'])} |
| CV MAE Std | {rupiah(metrics['cv_mae_std'])} |

Nilai evaluasi menunjukkan bahwa model dapat menangkap pola utama pada dataset dummy. Namun, hasil ini tidak boleh dianggap sebagai bukti performa pada bisnis nyata karena dataset masih sintetis. Nilai performa terutama menunjukkan bahwa pipeline, fitur, dan mekanisme prediksi sudah berjalan.

### 4.3 Feature Importance

| Fitur | Importance |
|---|---:|
{importance_rows}

Feature importance membantu menjelaskan faktor yang paling berpengaruh pada prediksi. Pada data sintetis ini, stok awal, harga, produk, dan kategori biasanya menjadi faktor penting karena generator dataset memang mendesain hubungan bisnis tersebut.

### 4.4 Aplikasi Streamlit

Aplikasi Streamlit dibuat dalam tiga halaman: Dashboard Data, Prediksi Penjualan, dan Evaluasi Model. Halaman prediksi menerima input produk, tanggal, harga, stok, promo, kanal, cuaca, pengunjung, dan rating, lalu menampilkan prediksi total penjualan dalam rupiah.

## BAB V. Kesimpulan dan Saran

### 5.1 Kesimpulan

Prototype sistem prediksi penjualan UMKM menggunakan Random Forest berhasil dibuat. Sistem mencakup dataset dummy yang aman, notebook eksperimen, pipeline model, aplikasi Streamlit multipage, dan laporan akademik. Random Forest cocok untuk prototype ini karena kuat pada data tabular dan interpretasinya relatif mudah.

### 5.2 Keterbatasan

Keterbatasan utama adalah penggunaan dataset sintetis. Pola yang muncul berasal dari desain generator, bukan observasi lapangan. Model juga belum diuji pada perubahan harga nyata, event lokal, gangguan stok, perubahan perilaku pelanggan, atau data multi-cabang.

### 5.3 Saran Pengembangan

Tahap profesional sebaiknya mengganti dataset dummy dengan data transaksi POS asli, menambahkan validasi kualitas data, menyimpan prediksi harian, membandingkan beberapa algoritma, melakukan retraining berkala, dan memonitor error prediksi setelah dipakai oleh UMKM.

## Daftar Pustaka

Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32. https://doi.org/10.1023/A:1010933404324

Punia, S., Nikolopoulos, K., Singh, S. P., Madaan, J. K., & Litsiou, K. (2020). Deep learning with long short-term memory networks and random forests for demand forecasting in multi-channel retail. International Journal of Production Research. https://doi.org/10.1080/00207543.2020.1735666

scikit-learn Developers. (2026). RandomForestRegressor and Pipeline documentation. https://scikit-learn.org/stable/

Streamlit. (2026). Multipage apps and caching documentation. https://docs.streamlit.io/

## Pernyataan Penggunaan AI

Dokumen, dataset dummy, kode, dan struktur prototype disusun dengan bantuan AI coding assistant. Pemeriksaan akhir, penyesuaian konteks akademik, dan validasi hasil tetap perlu dilakukan oleh penulis proyek sebelum dikumpulkan.
"""


def build_docx(markdown_text: str) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("| "):
            document.add_paragraph(line)
        elif line[0:3] in {"1. ", "2. ", "3. ", "4. ", "5. "}:
            document.add_paragraph(line[3:], style="List Number")
        else:
            document.add_paragraph(line)

    REPORT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(REPORT_DOCX)


def main() -> None:
    markdown_text = build_markdown()
    REPORT_MD.parent.mkdir(parents=True, exist_ok=True)
    REPORT_MD.write_text(markdown_text, encoding="utf-8")
    build_docx(markdown_text)
    print(f"saved {REPORT_MD} and {REPORT_DOCX}")


if __name__ == "__main__":
    main()
